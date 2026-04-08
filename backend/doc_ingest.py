# ============================================================
# FILE: backend/doc_ingest.py
# PURPOSE: Parse uploaded documents (PDF, DOCX, TXT, MD, etc.),
#          chunk them, embed with SentenceTransformers, and store
#          in a separate FAISS index for document-based RAG.
# ============================================================

import os
import json
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
from tqdm import tqdm


# ── Configuration ────────────────────────────────────────────
DOC_INDEX_DIR   = "doc_faiss_index"     # separate from code index
CHUNK_CHARS     = 1500                   # ~300 words per chunk
OVERLAP_CHARS   = 200                    # character overlap
BATCH_SIZE      = 32
EMBED_MODEL     = "sentence-transformers/all-MiniLM-L6-v2"
UPLOAD_DIR      = "uploads"              # where uploaded files are saved

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".rtf"}


def ensure_dirs():
    """Create necessary directories."""
    os.makedirs(DOC_INDEX_DIR, exist_ok=True)
    os.makedirs(UPLOAD_DIR, exist_ok=True)


# ── Parsers ──────────────────────────────────────────────────

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        return "\n\n".join(text_parts)
    except ImportError:
        raise ImportError(
            "PyPDF2 is required for PDF parsing. "
            "Install it with: pip install PyPDF2"
        )


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install it with: pip install python-docx"
        )


def parse_txt(file_path: str) -> str:
    """Read a plain text / markdown / CSV file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Could not read file: {e}")


def parse_document(file_path: str) -> str:
    """
    Detect file type and extract text content.

    Args:
        file_path: Path to the uploaded document.

    Returns:
        Extracted text content as a string.

    Raises:
        ValueError: If the file type is not supported.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in {".docx", ".doc"}:
        return parse_docx(file_path)
    elif ext in {".txt", ".md", ".csv", ".rtf"}:
        return parse_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


# ── Chunking ─────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = CHUNK_CHARS,
               overlap: int = OVERLAP_CHARS) -> list[str]:
    """
    Split document text into overlapping character-based chunks.

    Uses paragraph-aware splitting — tries to break at paragraph
    boundaries when possible, falling back to sentence boundaries,
    then hard character splits.

    Args:
        text:       Full document text.
        chunk_size: Target characters per chunk.
        overlap:    Characters of overlap between chunks.

    Returns:
        List of chunk strings.
    """
    if not text.strip():
        return []

    # If text is small enough, return as one chunk
    if len(text) <= chunk_size:
        return [text.strip()]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size

        # Try to break at a paragraph boundary
        if end < len(text):
            # Look for paragraph break near the end
            para_break = text.rfind("\n\n", start + chunk_size // 2, end + 100)
            if para_break != -1:
                end = para_break
            else:
                # Try sentence boundary (. or \n)
                sent_break = text.rfind(". ", start + chunk_size // 2, end + 50)
                if sent_break != -1:
                    end = sent_break + 1
                else:
                    newline_break = text.rfind("\n", start + chunk_size // 2, end + 50)
                    if newline_break != -1:
                        end = newline_break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap
        if start >= len(text):
            break

    return chunks


# ── Index Building ───────────────────────────────────────────

def build_doc_index(file_paths: list[str], file_names: list[str],
                    index_dir: str = DOC_INDEX_DIR) -> int:
    """
    Main pipeline: parse documents → chunk → embed → build FAISS index.

    Args:
        file_paths: List of absolute paths to uploaded files.
        file_names: List of original file names (for metadata).
        index_dir:  Directory to save FAISS index and metadata.

    Returns:
        Total number of chunks indexed.
    """
    os.makedirs(index_dir, exist_ok=True)

    all_chunks = []
    all_metadata = []

    for file_path, file_name in zip(file_paths, file_names):
        print(f"\nParsing: {file_name}")

        try:
            content = parse_document(file_path)
        except Exception as e:
            print(f"  WARNING: Failed to parse '{file_name}': {e}")
            continue

        if not content.strip():
            print(f"  WARNING: No text extracted from '{file_name}'")
            continue

        chunks = chunk_text(content)
        print(f"  Extracted {len(content)} chars → {len(chunks)} chunks")

        for i, chunk in enumerate(chunks):
            all_chunks.append(chunk)
            all_metadata.append({
                "chunk_id": len(all_metadata),
                "file_name": file_name,
                "chunk_index": i,
                "chunk_text": chunk,
            })

    print(f"\nTotal chunks to embed: {len(all_chunks)}")

    if not all_chunks:
        return 0

    # ── Embed all chunks ──────────────────────────────────────
    print(f"Loading embedding model: {EMBED_MODEL}")
    model = SentenceTransformer(EMBED_MODEL)

    print("Embedding chunks...")
    embeddings = model.encode(
        all_chunks,
        batch_size=BATCH_SIZE,
        show_progress_bar=True,
        convert_to_numpy=True,
        normalize_embeddings=True,
    )
    embeddings = embeddings.astype(np.float32)

    # ── Build FAISS index ─────────────────────────────────────
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    faiss.normalize_L2(embeddings)
    index.add(embeddings)

    # ── Save ──────────────────────────────────────────────────
    index_path = os.path.join(index_dir, "index.faiss")
    metadata_path = os.path.join(index_dir, "metadata.json")

    faiss.write_index(index, index_path)
    with open(metadata_path, "w", encoding="utf-8") as f:
        json.dump(all_metadata, f, ensure_ascii=False, indent=2)

    print(f"  FAISS index saved: {index_path}")
    print(f"  Metadata saved:    {metadata_path}")
    print(f"  Total vectors:     {index.ntotal}")

    return index.ntotal
